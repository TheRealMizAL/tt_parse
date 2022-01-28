from .abc import AbstractDocument
from io import BytesIO
import docx
import docx2txt
import os


class WordDocument(AbstractDocument):

    def __init__(self):
        self.__IMAGES_PATH = './images'
        self.__changes = {'group_name': 'ТО-21-1',
                          'changelist': [],
                          'images': []}

        try:
            os.mkdir(self.__IMAGES_PATH)
        except FileExistsError:
            pass

    async def parse(self, file_bytes):

        docx_file = docx.Document(BytesIO(file_bytes))
        await self.__tables_parse(docx_file.tables)

        docx2txt.process(BytesIO(file_bytes), self.__IMAGES_PATH)
        for image in os.listdir(self.__IMAGES_PATH):
            with open(f'{self.__IMAGES_PATH}/{image}', 'r') as img:
                self.__changes['images'].append(img)

        print(self.__changes)
        return self.__changes

    async def __tables_parse(self, tables):
        for table in tables:
            for row in table.rows[1:]:
                if row.cells[0] == self.__changes['group_name']:
                    lesson_name = row.cells[2].text.split('(')[0] + ' ' + row.cells[3].text
                    try:
                        teacher_name = row.cells[2].text.split('(')[1][:-1]
                    except IndexError:
                        teacher_name = None
                    for lesson_number in row.cells[1].text.split(','):
                        self.__changes['changelist'].append({'number': int(lesson_number) if lesson_number != '' else None,
                                                             'lesson_name': lesson_name,
                                                             'teacher_name': teacher_name})
